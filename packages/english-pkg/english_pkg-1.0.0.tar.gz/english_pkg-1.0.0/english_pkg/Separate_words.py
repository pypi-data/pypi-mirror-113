import docx
from docx.shared import Inches, Pt, RGBColor


class Separate_english(object):
    '''
    open(self,url): 打开一个需要处理的文件，url为需要文件的路径，返回值为该文件对象
    nrows(self):文件对象调用nrow（），返回打开文件的段落数
    count_of_type(self): 计算打开文件的单词词性分类，
    write(self):根据统计的分词种类，写入文档
    save(self,name='自动保存文件.docx'):将缓存中的内容写入文件name中，永久保存

    '''
    def __init__(self):
        # self.file = None
        self.document_write = docx.Document()
        pass

    def open(self,url):
        self.file = docx.Document(r"{}".format(url))
        return self.file

    def nrows(self):
        nrow = '段落: ' + str(len(self.file.paragraphs))
        return nrow

    def count_of_type(self):
        # 声明列表list，用于存储单词及其释意，
        self.list = []
        cixinglist = {}
        for i in range(len(self.file.paragraphs)):
            if len(self.file.paragraphs[i].text) == 2:
                continue
            # print(file.paragraphs[i].text.split(' '))

            # 将单词按空格分割后存储在列表中，方便后续程序代码，提出词性
            self.list.append(self.file.paragraphs[i].text.split(' '))

        # 遍历list里的所有单词，对相应词性进行计数统计，方便后续排除提取词性过程中的误操作
        for i in self.list:
            for j in i:
                if j.endswith('.'):
                    if j in cixinglist:
                        cixinglist[j] += 1
                    else:
                        cixinglist[j] = 1

        # 遍历list里的所有单词，过滤掉提取词性过程中的误提取，
        de = []
        for i in cixinglist:
            if cixinglist[i] < 2:
                de.append(i)

        # 删除误操作的词性
        for i in de:
            del cixinglist[i]

        for i in cixinglist:
            if i == 'adj.':
                cixinglist[i] = '形容词'
            elif i == 'n.':
                cixinglist[i] = '名词'
            elif i == 'adv.':
                cixinglist[i] = '副词'
            elif i == 'v.':
                cixinglist[i] = '动词'
            else:
                print(None)
        print(cixinglist)
        self.cixing = cixinglist

        return self.cixing

    def write(self):
        # 遍历词性列表里的词性，去打开的文件遍历，将对应词性的单词写入文档
        for ci in self.cixing:
            # 标题写入词性，并且加粗
            p1 = self.document_write.add_heading(u'{}:{}'.format(ci.split('.')[0],self.cixing[ci]), 1)
            p1.bold = True

            for h in self.list:
                # print(h)
                if ci in h:
                    para = '  '.join(h)
                    # print(para)
                    p5 = self.document_write.add_paragraph(para)

                    # 设置行间距和缩进
                    p5.paragraph_format.line_spacing = Pt(12)
                    p5.paragraph_format.left_indent = Inches(0.1)

            # 写完一个词性，换一页
            self.document_write.add_page_break()

    def save(self,name='自动保存文件.docx'):
        self.document_write.save(name)